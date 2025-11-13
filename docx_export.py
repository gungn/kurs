from docx import Document
from datetime import datetime

def save_calculation_to_docx(shape, parameters, perimeter, area):
    """
    Save calculation results to DOCX file
    """
    doc = Document()
    
    # Add title
    doc.add_heading('Geometry Calculation Results', 0)
    
    # Add calculation details
    doc.add_paragraph(f'Calculation date: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph(f'Shape: {shape}')
    
    # Add parameters
    params_text = ', '.join([f'{key} = {value}' for key, value in parameters.items()])
    doc.add_paragraph(f'Parameters: {params_text}')
    
    # Add results
    doc.add_paragraph(f'Perimeter: {perimeter:.2f}')
    doc.add_paragraph(f'Area: {area:.2f}')
    
    # Save file
    filename = f"geometry_calculation_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(filename)
    
    return filename